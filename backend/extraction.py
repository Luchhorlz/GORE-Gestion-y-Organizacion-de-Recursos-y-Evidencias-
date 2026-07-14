from __future__ import annotations

import io
import re
import threading
from dataclasses import dataclass
from pathlib import Path


SUPPORTED_MEDIA_TYPES = {
    "text/plain",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "image/jpeg",
    "image/png",
    "image/gif",
    "image/webp",
}
MAX_EXTRACTED_CHARACTERS = 5_000_000
MAX_PDF_PAGES = 500
OCR_LOCK = threading.Lock()
OCR_ENGINE = None


@dataclass(frozen=True)
class ExtractedSection:
    section_type: str
    section_label: str
    section_index: int
    text: str
    method: str


def normalize_text(value: str) -> str:
    value = value.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _get_ocr_engine():
    global OCR_ENGINE
    with OCR_LOCK:
        if OCR_ENGINE is None:
            from rapidocr_onnxruntime import RapidOCR

            OCR_ENGINE = RapidOCR()
        return OCR_ENGINE


def _ocr_image(image) -> str:
    import numpy as np

    engine = _get_ocr_engine()
    result, _elapsed = engine(np.asarray(image.convert("RGB")))
    if not result:
        return ""
    return normalize_text("\n".join(str(line[1]) for line in result if len(line) > 1 and line[1]))


def _extract_image(path: Path) -> list[ExtractedSection]:
    from PIL import Image

    with Image.open(path) as image:
        text = _ocr_image(image)
    return [ExtractedSection("image", "Imagen", 1, text, "rapidocr")]


def _extract_txt(path: Path) -> list[ExtractedSection]:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16", "cp1252"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")
    return [ExtractedSection("document", "Documento", 1, normalize_text(text), "plain_text")]


def _extract_docx(path: Path) -> list[ExtractedSection]:
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return [ExtractedSection("document", "Documento", 1, normalize_text("\n".join(parts)), "python-docx")]


def _extract_xlsx(path: Path) -> list[ExtractedSection]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[ExtractedSection] = []
    try:
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() if value is not None else "" for value in row]
                if any(values):
                    lines.append(" | ".join(values))
            sections.append(ExtractedSection("sheet", sheet.title, index, normalize_text("\n".join(lines)), "openpyxl"))
    finally:
        workbook.close()
    return sections


def _extract_pdf(path: Path) -> list[ExtractedSection]:
    from pypdf import PdfReader
    import pypdfium2 as pdfium

    reader = PdfReader(path)
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError("pdf_page_limit")
    pdf = pdfium.PdfDocument(path)
    sections: list[ExtractedSection] = []
    try:
        for index, page in enumerate(reader.pages, start=1):
            native = normalize_text(page.extract_text() or "")
            method = "pypdf"
            text = native
            if len(native) < 40:
                rendered = pdf[index - 1].render(scale=2).to_pil()
                ocr_text = _ocr_image(rendered)
                if ocr_text:
                    text = normalize_text("\n".join(part for part in (native, ocr_text) if part))
                    method = "pypdf+rapidocr" if native else "rapidocr"
            sections.append(ExtractedSection("page", f"Página {index}", index, text, method))
    finally:
        pdf.close()
    return sections


def extract_document(path: Path, media_type: str) -> list[ExtractedSection]:
    if media_type not in SUPPORTED_MEDIA_TYPES:
        return []
    if media_type == "text/plain":
        sections = _extract_txt(path)
    elif media_type == "application/pdf":
        sections = _extract_pdf(path)
    elif media_type.endswith("wordprocessingml.document"):
        sections = _extract_docx(path)
    elif media_type.endswith("spreadsheetml.sheet"):
        sections = _extract_xlsx(path)
    elif media_type.startswith("image/"):
        sections = _extract_image(path)
    else:
        sections = []
    total = sum(len(section.text) for section in sections)
    if total > MAX_EXTRACTED_CHARACTERS:
        raise ValueError("extracted_text_limit")
    return sections


def chunk_text(text: str, target_size: int = 1200, overlap: int = 180) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + target_size)
        if end < len(text):
            boundary = max(text.rfind("\n", start + target_size // 2, end), text.rfind(". ", start + target_size // 2, end))
            if boundary > start:
                end = boundary + 1
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return [chunk for chunk in chunks if chunk]
